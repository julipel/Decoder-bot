"""Тесты `DocxParser` (Sprint 6, задача S6-05, §14.2/14.11)."""

from __future__ import annotations

import io

import docx
import pytest

from dekoder.infrastructure.documents.parsers.docx_parser import DocxParser
from dekoder.shared.errors import ValidationError


def _build_docx(*, with_content: bool = True) -> bytes:
    document = docx.Document()
    if with_content:
        document.add_heading("Заголовок раздела", level=1)
        document.add_paragraph("Первый абзац.")
        document.add_heading("Подраздел", level=2)
        document.add_paragraph("Второй абзац.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class TestDocxParser:
    def test_normalizes_headings_to_markdown_and_keeps_paragraphs(self) -> None:
        parser = DocxParser()

        pages = parser.parse(_build_docx())

        assert len(pages) == 1
        assert pages[0].number is None
        assert pages[0].text == "# Заголовок раздела\n\nПервый абзац.\n\n## Подраздел\n\nВторой абзац."

    def test_title_style_becomes_a_level_one_heading(self) -> None:
        document = docx.Document()
        document.add_heading("Заголовок документа", level=0)  # python-docx: level=0 -> стиль "Title"
        document.add_paragraph("Текст.")
        buffer = io.BytesIO()
        document.save(buffer)
        parser = DocxParser()

        pages = parser.parse(buffer.getvalue())

        assert pages[0].text == "# Заголовок документа\n\nТекст."

    def test_blank_paragraphs_are_skipped(self) -> None:
        document = docx.Document()
        document.add_paragraph("Первый абзац.")
        document.add_paragraph("")
        document.add_paragraph("Второй абзац.")
        buffer = io.BytesIO()
        document.save(buffer)
        parser = DocxParser()

        pages = parser.parse(buffer.getvalue())

        assert pages[0].text == "Первый абзац.\n\nВторой абзац."

    def test_empty_document_raises_validation_error(self) -> None:
        parser = DocxParser()

        with pytest.raises(ValidationError) as exc_info:
            parser.parse(_build_docx(with_content=False))

        assert exc_info.value.code == "KNOWLEDGE_DOCUMENT_UNSUPPORTED"

    def test_corrupted_content_raises_validation_error(self) -> None:
        parser = DocxParser()

        with pytest.raises(ValidationError) as exc_info:
            parser.parse(b"not a docx file")

        assert exc_info.value.code == "KNOWLEDGE_DOCUMENT_UNSUPPORTED"
